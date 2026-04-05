import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Extract raw text from PDF, DOCX, or image bytes.
    Returns cleaned text string.
    """
    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "image": _extract_image,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    text = extractor(file_bytes)
    return _clean_text(text)


def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF.
    If the PDF is image-based (no text layer), automatically falls back
    to rendering each page as a high-res image and running Tesseract OCR.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        # --- Pass 1: Try native text extraction ---
        text_parts = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            t = page.get_text("text").strip()
            if t:
                text_parts.append(t)

        if text_parts:
            doc.close()
            logger.info(f"PDF: native text extracted from {len(text_parts)} page(s)")
            return "\n".join(text_parts)

        # --- Pass 2: OCR fallback for image-based / scanned PDFs ---
        logger.info("PDF has no text layer — falling back to OCR on rendered pages")
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            doc.close()
            raise RuntimeError("pytesseract/Pillow required for image-based PDF OCR")

        ocr_parts = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            # Render at 2x resolution for better OCR accuracy
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png")))

            try:
                text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
            except Exception:
                text = pytesseract.image_to_string(img, config="--psm 6")

            if text.strip():
                ocr_parts.append(text.strip())
                logger.info(f"PDF OCR page {page_num + 1}: {len(text.strip())} chars")

        doc.close()
        return "\n\n".join(ocr_parts)

    except RuntimeError:
        raise
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise RuntimeError(f"Failed to extract PDF text: {e}")


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also grab text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    table_text.append(" | ".join(row_data))

        return "\n".join(paragraphs + table_text)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise RuntimeError(f"Failed to extract DOCX text: {e}")


def _extract_image(file_bytes: bytes) -> str:
    """Extract text from image using Tesseract OCR."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise RuntimeError("pytesseract or Pillow not installed.")

    try:
        image = Image.open(io.BytesIO(file_bytes))

        # Try English first, fallback to multi-language
        try:
            text = pytesseract.image_to_string(image, lang="eng+hin", config="--psm 6")
        except Exception:
            text = pytesseract.image_to_string(image, config="--psm 6")

        return text
    except Exception as e:
        logger.error(f"OCR extraction error: {e}")
        raise RuntimeError(f"Failed to extract text from image: {e}")


def _clean_text(text: str) -> str:
    """Remove excessive whitespace and normalize text."""
    if not text:
        return ""
    # Remove excessive blank lines
    lines = [line.strip() for line in text.splitlines()]
    non_empty = [line for line in lines if line]
    return "\n".join(non_empty)
